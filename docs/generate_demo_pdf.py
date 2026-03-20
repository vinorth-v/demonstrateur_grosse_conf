"""Génère DEMO_finale.docx — version compacte pour impression."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import re

# Couleurs
BLEU_FONCE = RGBColor(0x1A, 0x3C, 0x6E)
BLEU_MOYEN = RGBColor(0x2D, 0x5F, 0xA1)
GRIS_FONCE = RGBColor(0x33, 0x33, 0x33)
GRIS_MOYEN = RGBColor(0x55, 0x55, 0x55)
ORANGE = RGBColor(0xE8, 0x6C, 0x00)
VERT = RGBColor(0x1B, 0x7A, 0x3D)
ROUGE = RGBColor(0xC0, 0x39, 0x2B)

BODY_SIZE = Pt(8.5)
SMALL_SIZE = Pt(7.5)

doc = Document()

# -- Marges serrées --
for section in doc.sections:
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

# -- Style de base compact --
style = doc.styles["Normal"]
style.font.name = "Helvetica"
style.font.size = BODY_SIZE
style.font.color.rgb = GRIS_FONCE
style.paragraph_format.space_after = Pt(1)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.line_spacing = 1.05


def set_cell_shading(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn("w:shd"), {
        qn("w:fill"): color_hex, qn("w:val"): "clear",
    })
    shading.append(shading_elem)


def set_cell_margins(cell, top=40, bottom=40, left=80, right=80):
    tc = cell._element.get_or_add_tcPr()
    margins = tc.makeelement(qn("w:tcMar"), {})
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        m = margins.makeelement(qn(f"w:{side}"), {
            qn("w:w"): str(val), qn("w:type"): "dxa",
        })
        margins.append(m)
    tc.append(margins)


def add_title():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run("DEMO FINALE — KYC avec LLM Multimodal")
    run.font.size = Pt(16)
    run.font.color.rgb = BLEU_FONCE
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("5-6 min  |  RAD & LAD  |  Structured Output")
    run.font.size = Pt(8)
    run.font.color.rgb = GRIS_MOYEN

    # Ligne séparatrice
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("─" * 90)
    run.font.size = Pt(5)
    run.font.color.rgb = BLEU_MOYEN


def add_phase_header(num, title, timing):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)

    run = p.add_run(f"PHASE {num} — ")
    run.font.size = Pt(11)
    run.font.color.rgb = BLEU_MOYEN
    run.bold = True
    run = p.add_run(title)
    run.font.size = Pt(11)
    run.font.color.rgb = BLEU_FONCE
    run.bold = True
    run = p.add_run(f"  ({timing})")
    run.font.size = Pt(7.5)
    run.font.color.rgb = GRIS_MOYEN


def add_screen(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Cm(0.3)
    run = p.add_run(f"ÉCRAN  {text}")
    run.font.size = SMALL_SIZE
    run.font.color.rgb = ORANGE
    run.bold = True


def _add_rich_text(paragraph, text, size=BODY_SIZE, default_color=GRIS_FONCE):
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.size = size
            run.font.color.rgb = BLEU_FONCE
            run.font.name = "Helvetica"
        else:
            run = paragraph.add_run(part)
            run.font.size = size
            run.font.color.rgb = default_color
            run.font.name = "Helvetica"


def add_speech(text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "EEF2F8")
    set_cell_margins(cell, top=30, bottom=30, left=70, right=50)

    # Bordure gauche bleue
    tc = cell._element.get_or_add_tcPr()
    borders = tc.makeelement(qn("w:tcBorders"), {})
    left = borders.makeelement(qn("w:left"), {
        qn("w:val"): "single", qn("w:sz"): "10",
        qn("w:color"): "2D5FA1", qn("w:space"): "0",
    })
    borders.append(left)
    tc.append(borders)

    cell.paragraphs[0].clear()
    lines = text.strip().split("\n")
    for i, line in enumerate(lines):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.05
        if line.strip() == "":
            run = p.add_run(" ")
            run.font.size = Pt(4)
        else:
            _add_rich_text(p, line, size=BODY_SIZE)

    # Petit espace après
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(1)
    sp.paragraph_format.space_before = Pt(0)
    run = sp.add_run("")
    run.font.size = Pt(1)


def add_note(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run("→ ")
    run.font.color.rgb = VERT
    run.font.size = SMALL_SIZE
    run.bold = True
    run = p.add_run(text)
    run.font.size = SMALL_SIZE
    run.font.color.rgb = GRIS_MOYEN
    run.italic = True


def add_section_title(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(13)
    run.font.color.rgb = BLEU_FONCE
    run.bold = True
    # Ligne
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(2)
    run = p2.add_run("─" * 90)
    run.font.size = Pt(5)
    run.font.color.rgb = BLEU_MOYEN


def add_question(question, severity=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(1)

    colors = {"critical": ROUGE, "probable": ORANGE, "bonus": VERT}
    labels = {"critical": "CRITIQUE", "probable": "PROBABLE", "bonus": "BONUS"}
    if severity and severity in colors:
        run = p.add_run(f"[{labels[severity]}] ")
        run.font.size = Pt(7)
        run.font.color.rgb = colors[severity]
        run.bold = True

    run = p.add_run(f"« {question} »")
    run.font.size = Pt(9)
    run.font.color.rgb = BLEU_FONCE
    run.bold = True


def add_answer(text):
    add_speech(text)


def add_bullets(items):
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run("• ")
        run.font.color.rgb = BLEU_MOYEN
        run.font.size = BODY_SIZE
        _add_rich_text(p, item, size=BODY_SIZE)


def add_text(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(1)
    _add_rich_text(p, text, size=BODY_SIZE)


def add_code(text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F4F4")
    set_cell_margins(cell, top=25, bottom=25, left=60, right=60)
    cell.paragraphs[0].clear()
    for i, line in enumerate(text.strip().split("\n")):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run(line)
        run.font.size = Pt(7.5)
        run.font.name = "Courier New"
        run.font.color.rgb = GRIS_FONCE
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(1)


# ============================================================
# CONTENU
# ============================================================

add_title()

# --- PHASE 1 ---
add_phase_header(1, "La problématique", "0:00 — 0:45")
add_screen("Face au public, pas de slide")
add_speech(
    "Est-ce qu'il vous est déjà arrivé de devoir classifier et extraire des informations "
    "de milliers, voire de millions de documents ?\n"
    "C'est le problème de la **KYC — Know Your Customer**. Le client nous envoie des CNI, "
    "des passeports, des justificatifs de domicile, des RIB... et on doit extraire les "
    "informations en temps réel.\n"
    "Il y a deux grandes étapes :\n"
    "— La **RAD** — Reconnaissance Automatique de Documents : c'est quoi ce document ?\n"
    "— La **LAD** — Lecture Automatique de Documents : qu'est-ce qu'il y a dedans ?"
)
add_speech(
    "Avant, pour faire ça, il fallait : collecter un dataset, le faire annoter, reviewer les annotations, "
    "entraîner un modèle de Computer Vision, l'évaluer, le réentraîner...\n"
    "On parle de **6 mois de développement, 3 à 5 ML engineers**, des annotateurs, et facilement **100 000€**.\n"
    "Ce que je vais vous montrer, c'est que tout ça peut être remplacé par un **prompt** et un **schéma de données**."
)

# --- PHASE 2 ---
add_phase_header(2, "Le prompt, c'est le nouveau modèle", "0:45 — 1:45")
add_screen("src/chains/prompts.py — PROMPT_CLASSIFICATION")
add_speech(
    "Voici le prompt de classification — la RAD. On décrit simplement au LLM les types de documents "
    "qu'on attend, et on lui demande de nous dire lequel c'est, avec un score de confiance.\n"
    "**Trente lignes de texte.** C'est tout. Pas de CNN, pas de dataset, pas d'entraînement.\n"
    "Le LLM est multimodal : il **voit** l'image et il **comprend** ce qu'il voit."
)
add_screen("Scroller vers PROMPT_EXTRACTION_PASSEPORT")
add_speech(
    "Et voici le prompt d'extraction — la LAD. On lui dit : « extrais le nom, prénom, date de naissance, "
    "numéro, MRZ... » On guide le modèle avec des instructions précises : format des dates, lecture de la "
    "zone MRZ, champs obligatoires.\n"
    "C'est du **prompt engineering** — et c'est ça qui remplace des mois d'annotation et d'entraînement."
)

# --- PHASE 3 ---
add_phase_header(3, "Le structured output, la clé de voûte", "1:45 — 3:00")
add_screen("src/chains/schemas/kyc_schemas.py — classe Passeport")
add_speech(
    "Comment on **structure** tout ça correctement ? Comment on évite que le LLM renvoie n'importe quoi ?\n"
    "La réponse : le **structured output** avec **Pydantic**.\n"
    "Chaque champ a un **type** (str, date, Optional) et une **description** qui aide le LLM. "
    "Et Pydantic **valide** la sortie : mauvais format → erreur. Champ manquant → erreur. "
    "On ne fait **jamais confiance aveuglément** à une sortie de modèle."
)
add_screen("Classe RIB — validateur IBAN")
add_speech(
    "On ajoute des **règles métier** dans le schéma. Ici, un validateur vérifie le **checksum IBAN** — "
    "l'algorithme modulo 97. Si l'IBAN ne passe pas, on le sait immédiatement.\n"
    "C'est comme ça qu'on gère les hallucinations : le **LLM extrait**, les **règles métier valident**."
)
add_screen("Classe JustificatifDomicile — validateur de récence")
add_speech(
    "Autre exemple : le justificatif de domicile. On vérifie automatiquement qu'il date de "
    "**moins de 3 mois** — exigence réglementaire."
)

# --- PHASE 4 ---
add_phase_header(4, "La démo live", "3:00 — 4:30")
add_screen("examples/passeport.webp en plein écran")
add_speech(
    "Voici un passeport de test — des données fictives. Même pour un humain, lire la zone MRZ "
    "en bas, c'est fastidieux."
)
add_screen("Terminal → just passeport")
add_speech("Je lance le traitement.")
add_note("Pendant que ça tourne (~3-4 sec) :")
add_speech(
    "Le LLM reçoit l'image, la classifie (RAD), extrait les données (LAD), et Pydantic valide le tout."
)
add_note("Résultat affiché :")
add_speech(
    "**Classification** : passeport, score de confiance. Puis toutes les données : nom, prénom, "
    "date de naissance, numéro, MRZ complète... Tout structuré, tout validé.\n"
    "Et le coût : quelques fractions de centimes par document."
)
add_screen("2e terminal — résultat CNI pré-lancé")
add_speech(
    "Même chose pour une carte d'identité. Différent format, différentes données, mais "
    "**aucun changement de code, aucun réentraînement**. Le LLM s'adapte."
)

# --- PHASE 5 ---
add_phase_header(5, "L'impact et les résultats", "4:30 — 5:15")
add_speech(
    "Récapitulons :\n"
    "— **5 types de documents** : CNI, passeport, permis, justificatif domicile, RIB\n"
    "— **800 lignes de Python** (vs 10 000-15 000 en deep learning classique)\n"
    "— **Développement : 2 jours** (vs 6 mois)\n\n"
    "Sur un autre projet client (documents fiscaux), sur **6 000 documents** : "
    "**100% en classification**, **98.7% en extraction** — avec des cases à cocher.\n"
    "Les cases à cocher, c'est le cas qui tue. Permis 14 catégories ? Deep learning : 40-60h de dev. "
    "LLM multimodal : **15 minutes**."
)

# --- PHASE 6 ---
add_phase_header(6, "Le message à retenir", "5:15 — 5:45")
add_speech(
    "Données confidentielles → **environnement sécurisé**. Ici Vertex AI en Europe, "
    "mais l'architecture est découplée du provider.\n\n"
    "Pour faire de la RAD/LAD aujourd'hui, **trois choses** :\n"
    "1. Un **prompt** bien écrit qui guide le modèle\n"
    "2. Un **schéma Pydantic** qui structure et valide la sortie\n"
    "3. Des **règles métier** qui contrôlent les hallucinations\n\n"
    "On passe de mois d'entraînement à des jours de développement.\n"
    "Voulez-vous que je teste un autre document, ou on passe aux questions ?"
)

# ============================================================
# PRÉPARATION
# ============================================================

add_section_title("Préparation avant la démo")

add_text("**Commandes à préparer :**")
add_code("# Terminal 1 — live        # Terminal 2 — pré-lancé\njust passeport             just cni")

add_text("**Fichiers ouverts dans VS Code :**")
add_bullets([
    "**src/chains/prompts.py** — scrollé sur PROMPT_CLASSIFICATION",
    "**src/chains/schemas/kyc_schemas.py** — scrollé sur classe Passeport",
    "**examples/passeport.webp** (preview) + Terminal 1 prêt + Terminal 2 avec résultat CNI",
])

add_text("**Backup :**")
add_bullets([
    "Latence longue → « Pendant que ça tourne... » → montrer les schemas",
    "Vertex AI down → screenshots des résultats prêts",
])

add_text("**Tips :**")
add_bullets([
    "**Public mixte** : expliquer le code en termes simples",
    "**Impact** : insister sur les chiffres avant/après (6 mois → 2 jours, 100k€ → quelques €)",
])

# ============================================================
# FAQ
# ============================================================

add_section_title("FAQ — Réponses aux questions")

add_text("**Questions critiques**")

add_question("Comment vous gérez le RGPD / la confidentialité ?", "critical")
add_answer(
    "**Infrastructure** : Vertex AI en Europe (Frankfurt), données jamais hors UE, ISO 27001/SOC 2.\n"
    "**Contractuel** : DPA Google — aucune donnée n'entraîne leurs modèles. Zéro rétention.\n"
    "**Métier** : DPO a validé, DPIA positif. Chaque appel API loggé et auditable."
)

add_question("Et si Google change son pricing ou ses conditions ?", "critical")
add_answer(
    "**Architecture découplée** : basculer sur Claude, GPT-4V ou Mistral en ~20 lignes.\n"
    "**Exit strategy** : modèles open-source multimodaux progressent vite.\n"
    "**Pricing** : à 0.005€/doc, même x2 reste 100x moins cher que l'ancien process."
)

add_question("C'est en production ou un POC ?", "critical")
add_answer(
    "Le démonstrateur est un POC. Sur un autre périmètre, validé sur 6 000 documents réels. "
    "Objectif : industrialiser avec les briques de sécurité et monitoring."
)

add_text("**Questions probables**")

add_question("Comment vous gérez les hallucinations ?", "probable")
add_answer(
    "Le LLM extrait, les règles métier valident. Ex : checksum IBAN modulo 97. "
    "Échec → rejet automatique + revue humaine. Structured output Pydantic = filet de sécurité."
)

add_question("Documents de mauvaise qualité / manuscrits ?", "probable")
add_answer(
    "LLM plus robuste que les anciens OCR — utilise le contexte global. "
    "Seuil de confiance < 85% → revue humaine (~5% des docs). L'humain reste dans la boucle."
)

add_question("Pourquoi pas un modèle open-source on-premise ?", "probable")
add_answer(
    "Open-source multimodal = 6-12 mois de retard en compréhension visuelle. "
    "Self-hosting = GPU A100/H100, 50k€/an min. API = 2k€/an. "
    "Mais l'architecture est prête à basculer."
)

add_question("Ça coûte combien ?", "probable")
add_answer("~0.005€/document. Dossier KYC complet (identité + justificatif + RIB) : < 5 centimes.")

add_question("Temps de traitement ?", "probable")
add_answer("2-4 sec/document. Dossier KYC complet : ~10 sec. Ancien process humain : 48h minimum.")

add_question("Équipe nécessaire ?", "probable")
add_answer("Avant : 3-5 ML engineers + annotateurs × 6 mois. Maintenant : 1 dev × 2 jours.")

add_text("**Questions bonus**")

add_question("Qualité en production ?", "bonus")
add_answer(
    "3 KPIs : extraction réussie (cible 95%), rejet auto (< 5%), erreur humaine (< 1%). "
    "Échantillonnage 2% pour audit hebdo."
)

add_question("Explicabilité pour les auditeurs ?", "bonus")
add_answer("Prompt versionné dans Git. Règles métier explicites. Plus lisible qu'un CNN à 10M paramètres.")

add_question("Évolution des formats de documents ?", "bonus")
add_answer("Format CNI change → adapter le prompt/schéma. Pas de réannotation, pas de réentraînement.")

add_question("Documents dans d'autres langues ?", "bonus")
add_answer("LLMs nativement multilingues. Adapter le prompt suffit. Testé sur passeports étrangers sans changement.")

add_question("Erreur de classification ?", "bonus")
add_answer("Score de confiance. Sous le seuil → double vérification ou routage humain. En pratique : 100% en classif.")

add_question("Montée en charge ?", "bonus")
add_answer("Vertex AI scale auto. Parallélisation 10-100 docs simultanés. Pas de GPU à provisionner.")

# ============================================================
output_path = "DEMO_finale.docx"
doc.save(output_path)
print(f"✅ Document généré : {output_path}")
